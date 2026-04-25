"""Generate Efrain_Plascencia_Resume.docx from the site's resume-data.js,
matching the site's Swiss / Inter / JetBrains Mono aesthetic."""
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

# Colors pulled from styles.css (:root light theme). Accent is the
# sRGB conversion of oklch(0.42 0.12 255).
INK = RGBColor(0x1A, 0x1A, 0x1A)
INK_2 = RGBColor(0x3A, 0x3A, 0x3A)
INK_3 = RGBColor(0x6A, 0x6A, 0x6A)
INK_4 = RGBColor(0xA8, 0xA5, 0xA0)
ACCENT = RGBColor(0x15, 0x4C, 0x8C)
RULE = RGBColor(0xDD, 0xD9, 0xD1)

FF_DISPLAY = "Inter Tight"
FF_BODY = "Inter"
FF_MONO = "JetBrains Mono"

ROOT = Path(__file__).parent
OUTPUT = ROOT / "uploads" / "Efrain_Plascencia_Resume.docx"
MEMOJI = ROOT / "memoji_standard_transparent.png"
MEMOJI_CIRCLE = ROOT / "_memoji_circle.png"

# Site bg (--bg) and rule (--rule) from styles.css :root.
BG_COLOR = (0xF4, 0xF1, 0xEA, 255)
RULE_RGBA = (0xDD, 0xD9, 0xD1, 255)


def make_circular_memoji(src_path, out_path, size=600):
    """Bake the .floating-logo treatment around the memoji: circular --bg fill,
    1px --rule ring, memoji at 58/72 of the diameter (matching the site)."""
    src = Image.open(src_path).convert("RGBA")

    # Site ratios: 72px circle, 58px image, 1px border.
    inner_ratio = 58 / 72
    border_w = max(3, round(size / 72))  # ~1px scaled up; min 3 for crispness

    canvas = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(canvas)
    draw.ellipse([0, 0, size - 1, size - 1], fill=BG_COLOR)
    draw.ellipse([0, 0, size - 1, size - 1], outline=RULE_RGBA, width=border_w)

    inner = int(size * inner_ratio)
    resized = src.resize((inner, inner), Image.LANCZOS)
    offset = (size - inner) // 2
    canvas.alpha_composite(resized, (offset, offset))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(out_path, "PNG")
    return out_path


def load_data():
    text = (ROOT / "resume-data.js").read_text(encoding="utf-8")
    timeline_match = re.search(r"const TIMELINE\s*=\s*(\[.*?\]);", text, re.DOTALL)
    meta_match = re.search(r"const RESUME_META\s*=\s*(\{.*?\});", text, re.DOTALL)
    timeline = json.loads(timeline_match.group(1))
    meta = json.loads(meta_match.group(1))
    return timeline, meta


def set_character_spacing(run, points):
    """w:spacing is in 20ths of a point."""
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(int(points * 20)))
    rPr.append(spacing)


def set_font(run, name, size_pt, color=INK, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    # Also set east-asian and complex-script font so Word uses our font everywhere
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def add_bottom_border(paragraph, color=RULE, size=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_label(doc, number, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    num = p.add_run(f"{number}   ")
    set_font(num, FF_MONO, 9, INK_4)
    set_character_spacing(num, 1.0)
    lbl = p.add_run(label.upper())
    set_font(lbl, FF_MONO, 9, INK)
    set_character_spacing(lbl, 1.5)
    add_bottom_border(p, RULE)
    return p


def mono_line(doc, text, color=INK_3, size=8.5, space=1.0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, FF_MONO, size, color)
    set_character_spacing(r, space)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    dash = p.add_run("— ")
    set_font(dash, FF_BODY, 10, ACCENT)
    body = p.add_run(text)
    set_font(body, FF_BODY, 10, INK_2)
    return p


def tab_right(paragraph, indent_inches=7.5):
    """Add a right-aligned tab stop at indent_inches."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(indent_inches * 1440)))  # 1440 twips per inch
    tabs.append(tab)
    pPr.append(tabs)


def build_document(timeline, meta):
    doc = Document()

    # Thin margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Default paragraph style
    normal = doc.styles["Normal"]
    normal.font.name = FF_BODY
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK_2

    # --- HEADER: memoji + name + subtitle + contact ---
    circular = make_circular_memoji(MEMOJI, MEMOJI_CIRCLE)
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_after = Pt(4)
    img_p.add_run().add_picture(str(circular), width=Inches(1.1))

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(meta["name"])
    set_font(name_run, FF_DISPLAY, 32, INK)
    # tighter tracking like the site
    set_character_spacing(name_run, -0.6)

    subtitle_text = meta["titleLine"].upper().replace(" · ", "   ·   ")
    mono_line(doc, subtitle_text, INK_3, size=8.5, space=1.3, after=8)

    contact_bits = [
        meta.get("location", ""),
        meta.get("phone", ""),
        meta.get("email", ""),
        meta.get("linkedin", ""),
        meta.get("github", ""),
    ]
    contact_text = "   ·   ".join(b for b in contact_bits if b).upper()
    contact_p = mono_line(doc, contact_text, INK_2, size=8, space=1.0, after=2)
    add_bottom_border(contact_p, RULE)

    # --- 01 PROFESSIONAL SUMMARY ---
    section_label(doc, "01", "Professional Summary")
    p = doc.add_paragraph(meta["summary"])
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        set_font(r, FF_BODY, 10, INK_2)
    p.paragraph_format.line_spacing = 1.35

    # --- 02 CORE COMPETENCIES ---
    section_label(doc, "02", "Core Competencies")
    comps = meta.get("competencies", [])
    rows = (len(comps) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.autofit = True
    for i, comp in enumerate(comps):
        cell = table.cell(i // 2, i % 2)
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(2)
        num = cp.add_run(f"{i+1:02d}   ")
        set_font(num, FF_MONO, 8, INK_4)
        set_character_spacing(num, 0.8)
        body = cp.add_run(comp)
        set_font(body, FF_BODY, 9.5, INK_2)

    # --- 03 PROFESSIONAL EXPERIENCE & EDUCATION ---
    section_label(doc, "03", "Professional Experience & Education")
    for item in timeline:
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(8)
        head.paragraph_format.space_after = Pt(1)
        tab_right(head, 7.5)
        kind = "EDUCATION" if item.get("type") == "education" else "EXPERIENCE"
        kind_run = head.add_run(kind + "   ")
        set_font(kind_run, FF_MONO, 8, ACCENT)
        set_character_spacing(kind_run, 1.2)
        title_run = head.add_run(item["title"])
        set_font(title_run, FF_DISPLAY, 13, INK, bold=False)
        sep_run = head.add_run("\t")
        dates_run = head.add_run(item["dates"].upper())
        set_font(dates_run, FF_MONO, 8.5, INK_3)
        set_character_spacing(dates_run, 1.0)

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(4)
        org_run = meta_p.add_run(item["org"])
        set_font(org_run, FF_MONO, 9, INK_2)
        set_character_spacing(org_run, 0.6)
        dot_run = meta_p.add_run("   ·   ")
        set_font(dot_run, FF_MONO, 9, INK_4)
        city_run = meta_p.add_run(item["city"])
        set_font(city_run, FF_MONO, 9, INK_3)

        for bullet in item.get("bullets", []):
            add_bullet(doc, bullet)

    # --- 04 TECHNICAL SKILLS & TOOLS ---
    section_label(doc, "04", "Technical Skills & Tools")
    for category, desc in meta.get("skills", {}).items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        head_run = p.add_run(category.upper() + "   ")
        set_font(head_run, FF_MONO, 8.5, INK)
        set_character_spacing(head_run, 1.2)
        body = doc.add_paragraph(desc)
        body.paragraph_format.space_after = Pt(4)
        body.paragraph_format.left_indent = Inches(0.0)
        for r in body.runs:
            set_font(r, FF_BODY, 10, INK_2)
        body.paragraph_format.line_spacing = 1.3

    # --- 05 HIGHLIGHTS ---
    highlights = meta.get("highlights", [])
    if highlights:
        section_label(doc, "05", "Highlights")
        for h in highlights:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(4)
            head.paragraph_format.space_after = Pt(1)
            tab_right(head, 7.5)
            title_run = head.add_run(h["title"])
            set_font(title_run, FF_DISPLAY, 12, INK)
            head.add_run("\t")
            year_run = head.add_run(str(h["year"]))
            set_font(year_run, FF_MONO, 9, ACCENT)
            set_character_spacing(year_run, 1.0)
            body = doc.add_paragraph(h["body"])
            body.paragraph_format.space_after = Pt(4)
            for r in body.runs:
                set_font(r, FF_BODY, 10, INK_2)
            body.paragraph_format.line_spacing = 1.3

    # --- 06 PROJECTS placeholder (matches site) ---
    section_label(doc, "06", "Projects")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("COMING SOON")
    set_font(r, FF_MONO, 9, INK_3)
    set_character_spacing(r, 1.4)

    # --- 07 CERTIFICATIONS ---
    certs = meta.get("certifications", [])
    if certs:
        section_label(doc, "07", "Certifications")
        for c in certs:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(3)
            head.paragraph_format.space_after = Pt(2)
            tab_right(head, 7.5)
            title_run = head.add_run(c["title"])
            set_font(title_run, FF_DISPLAY, 12, INK)
            head.add_run("\t")
            year_run = head.add_run(str(c["year"]))
            set_font(year_run, FF_MONO, 9, ACCENT)
            set_character_spacing(year_run, 1.0)
            org = doc.add_paragraph()
            org.paragraph_format.space_after = Pt(2)
            org_run = org.add_run(c["org"])
            set_font(org_run, FF_MONO, 9, INK_3)
            set_character_spacing(org_run, 0.8)

    # --- 08 LANGUAGES ---
    languages = meta.get("languages", [])
    if languages:
        section_label(doc, "08", "Languages")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        for i, lang in enumerate(languages):
            if i > 0:
                sep = p.add_run("   ·   ")
                set_font(sep, FF_MONO, 9, INK_4)
            # German italicized per the site touch
            italic = lang["name"] == "German"
            name_run = p.add_run(lang["name"])
            set_font(name_run, FF_DISPLAY, 11, INK, italic=italic)
            level_run = p.add_run(" " + lang["level"].upper())
            set_font(level_run, FF_MONO, 8, INK_3)
            set_character_spacing(level_run, 1.0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    timeline, meta = load_data()
    build_document(timeline, meta)
